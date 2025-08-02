from docx import Document

def parse_document(doc_path):
    document = Document(doc_path)
    data = {"paragraphs": [], "tables": []}
    
    for para in document.paragraphs:
        data["paragraphs"].append({
            "text": para.text,
            "style": para.style.name
        })

    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        data["tables"].append(table_data)
    
    return data
